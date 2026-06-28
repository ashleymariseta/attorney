"""HTML → Word (.docx) for drafted legal documents.

Same firm standard as the PDF: **Book Antiqua 12pt, 1.5 line spacing, A4,
1-inch margins, justified body, numbered paragraphs with a hanging indent**,
and a **9pt "Title — Page N" footer** (a live Word PAGE field) on every page.

The body is HTML (or Markdown from older docs, normalised via
:func:`workflows.dochtml.to_html`). We walk it with lxml and emit python-docx
paragraphs/runs/tables.
"""
from __future__ import annotations

import base64
from io import BytesIO

import lxml.html as LH
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

FONT = 'Book Antiqua'
BODY_PT = 12
FOOTER_PT = 9

_ALIGN = {
    'left': WD_ALIGN_PARAGRAPH.LEFT, 'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'center': WD_ALIGN_PARAGRAPH.CENTER, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_INLINE_FMT = {
    'strong': 'bold', 'b': 'bold', 'em': 'italic', 'i': 'italic',
    'u': 'underline', 's': 'strike', 'del': 'strike', 'strike': 'strike', 'code': 'code',
}
_BLOCK_IN_P = {'img', 'ul', 'ol', 'table', 'blockquote', 'hr', 'pre', 'figure'}


def _tag(el) -> str:
    return (el.tag if isinstance(el.tag, str) else '').lower()


def _set_font(font, size_pt=None, bold=None, italic=None, underline=None, strike=None, mono=False):
    name = 'Courier New' if mono else FONT
    font.name = name
    # Ensure the east-asian/hAnsi slots also use the font (python-docx only sets ascii).
    rpr = font.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if underline is not None:
        font.underline = underline
    if strike is not None:
        font.strike = strike


def _add_runs(paragraph, node, fmt):
    """Recursively add formatted runs for an element's inline content."""
    def emit(text):
        if not text:
            return
        run = paragraph.add_run(text)
        _set_font(run.font, bold=fmt.get('bold'), italic=fmt.get('italic'),
                   underline=fmt.get('underline'), strike=fmt.get('strike'), mono=fmt.get('code'))
        if fmt.get('link'):
            run.font.color.rgb = RGBColor(0x1A, 0x57, 0xBF)

    emit(node.text)
    for child in node:
        t = _tag(child)
        if t == 'br':
            paragraph.add_run().add_break()
        elif t in _INLINE_FMT:
            sub = dict(fmt)
            sub[_INLINE_FMT[t]] = True
            _add_runs(paragraph, child, sub)
        elif t == 'a':
            sub = dict(fmt)
            sub['underline'] = True
            sub['link'] = True
            _add_runs(paragraph, child, sub)
        elif t in ('span', 'label', 'mark', 'p', 'div'):
            _add_runs(paragraph, child, fmt)
        elif t in _BLOCK_IN_P:
            pass  # block child handled by the block walker
        else:
            _add_runs(paragraph, child, fmt)
        if child.tail:
            r = paragraph.add_run(child.tail)
            _set_font(r.font, bold=fmt.get('bold'), italic=fmt.get('italic'),
                       underline=fmt.get('underline'), strike=fmt.get('strike'))


def _alignment(el):
    style = el.get('style', '')
    for key, val in _ALIGN.items():
        if 'text-align:%s' % key in style.replace(' ', '').lower():
            return val
    return None


def _add_picture(doc, src, max_w_in):
    if not src or not src.startswith('data:'):
        return
    try:
        raw = base64.b64decode(src.split(',', 1)[1])
        doc.add_picture(BytesIO(raw))
    except Exception:  # noqa: BLE001
        return
    pic = doc.paragraphs[-1]
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = doc.inline_shapes[-1] if doc.inline_shapes else None
    if shape and shape.width and shape.width > Inches(max_w_in):
        ratio = Inches(max_w_in) / shape.width
        shape.width = int(shape.width * ratio)
        shape.height = int(shape.height * ratio)


def _add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _walk(doc, el, max_w_in):
    t = _tag(el)
    if t in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if t == 'h1' else (_alignment(el) or WD_ALIGN_PARAGRAPH.LEFT)
        _add_runs(p, el, {'bold': True})
        for run in p.runs:
            _set_font(run.font, size_pt=BODY_PT, bold=True)
    elif t in ('p', 'div'):
        if (el.text_content() or '').strip():
            p = doc.add_paragraph()
            al = _alignment(el)
            if al is not None:
                p.alignment = al
            _add_runs(p, el, {})
        for child in el:
            if _tag(child) in _BLOCK_IN_P:
                _walk(doc, child, max_w_in)
    elif t == 'ul':
        for li in el:
            if _tag(li) != 'li':
                continue
            p = doc.add_paragraph(style='List Bullet')
            _add_runs(p, li, {})
            for sub in li:
                if _tag(sub) in ('ul', 'ol'):
                    _walk(doc, sub, max_w_in)
    elif t == 'ol':
        for li in el:
            if _tag(li) != 'li':
                continue
            p = doc.add_paragraph(style='List Number')
            _add_runs(p, li, {})
            for sub in li:
                if _tag(sub) in ('ul', 'ol'):
                    _walk(doc, sub, max_w_in)
    elif t == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        _add_runs(p, el, {'italic': True})
    elif t == 'hr':
        _add_hr(doc)
    elif t == 'pre':
        p = doc.add_paragraph()
        run = p.add_run(el.text_content() or '')
        _set_font(run.font, mono=True, size_pt=10)
    elif t == 'table':
        _add_table(doc, el)
    elif t == 'img':
        _add_picture(doc, el.get('src', ''), max_w_in)
    elif t in ('figure', 'section', 'article', 'header', 'footer', 'main'):
        for child in el:
            _walk(doc, child, max_w_in)
    else:
        if (el.text_content() or '').strip():
            p = doc.add_paragraph()
            _add_runs(p, el, {})


def _add_table(doc, table_el):
    rows = []
    for tr in table_el.iter('tr'):
        cells = [(cell, _tag(cell) == 'th') for cell in tr if _tag(cell) in ('td', 'th')]
        if cells:
            rows.append(cells)
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ''
            if ci < len(row):
                el, is_h = row[ci]
                p = cell.paragraphs[0]
                _add_runs(p, el, {'bold': True} if is_h else {})
                for run in p.runs:
                    _set_font(run.font, size_pt=11, bold=is_h or None)


def _add_page_field(paragraph, size_pt):
    """Footer run: a live Word PAGE field (so the number updates per page)."""
    run = paragraph.add_run()
    _set_font(run.font, size_pt=size_pt)
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def render_document_docx(title: str, body: str) -> bytes:
    from .dochtml import to_html

    title = (title or 'Document').strip()
    html = to_html(body)

    doc = Document()

    # Base "Normal" style: Arial 12pt, 1.5 spacing, justified.
    normal = doc.styles['Normal']
    _set_font(normal.font, size_pt=BODY_PT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Page geometry: A4, 1-inch margins.
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    max_w_in = 6.27  # A4 width − 2×1-inch margins

    try:
        root = LH.fragment_fromstring(html or '', create_parent='div')
    except Exception:  # noqa: BLE001
        root = LH.fromstring('<div>%s</div>' % (html or ''))
    if root.text and root.text.strip():
        doc.add_paragraph(root.text.strip())
    for el in root:
        _walk(doc, el, max_w_in)

    # 9pt "Title — Page N" footer with a live PAGE field.
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = footer_p.add_run(f'{title} — Page ')
    _set_font(lead.font, size_pt=FOOTER_PT)
    try:
        _add_page_field(footer_p, FOOTER_PT)
    except Exception:  # noqa: BLE001 — never let the footer break an export
        pass

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
